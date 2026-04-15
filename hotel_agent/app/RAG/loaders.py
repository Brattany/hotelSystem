from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class LoadedDocument:
    text: str
    metadata: Dict[str, Any]


def _load_txt_or_md(file_path: Path) -> LoadedDocument:
    text = file_path.read_text(encoding="utf-8")
    return LoadedDocument(
        text=text,
        metadata={
            "file_name": file_path.name,
            "source": file_path.name,
            "suffix": file_path.suffix.lower(),
        },
    )


def _load_pdf(file_path: Path) -> LoadedDocument:
    reader = PdfReader(str(file_path))
    pages = []

    for page in reader.pages:
        pages.append(page.extract_text() or "")

    text = "\n\n".join(pages)
    return LoadedDocument(
        text=text,
        metadata={
            "file_name": file_path.name,
            "source": file_path.name,
            "suffix": ".pdf",
            "page_count": len(reader.pages),
        },
    )


def _load_docx(file_path: Path) -> LoadedDocument:
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs)

    return LoadedDocument(
        text=text,
        metadata={
            "file_name": file_path.name,
            "source": file_path.name,
            "suffix": ".docx",
            "paragraph_count": len(paragraphs),
        },
    )


def load_document(file_path: str | Path) -> LoadedDocument:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return _load_txt_or_md(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)

    raise ValueError(f"暂不支持的文件类型: {suffix}")